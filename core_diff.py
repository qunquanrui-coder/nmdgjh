# -*- coding: utf-8 -*-
import bridge
import difflib
import threading
import time
import pandas as pd
from pathlib import Path
from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter

try:
    import docx
except Exception:
    docx = None


def _start_heartbeat(label, state, interval=5):
    stop_event = threading.Event()
    start_time = time.time()

    def worker():
        while not stop_event.wait(interval):
            stage = state.get("stage", "processing")
            current = state.get("current")
            total = state.get("total")
            elapsed = int(time.time() - start_time)
            if current is not None and total:
                bridge.update_terminal(f"HB [{label}] {stage}: {current}/{total}, elapsed {elapsed}s")
            else:
                bridge.update_terminal(f"HB [{label}] {stage}, elapsed {elapsed}s")

    thread = threading.Thread(target=worker, daemon=True)
    thread.start()
    return stop_event, thread


def _stop_heartbeat(stop_event, thread):
    stop_event.set()
    try:
        thread.join(timeout=1)
    except Exception:
        pass


def _read_docx_lines(path: Path, progress=None):
    d = docx.Document(str(path))
    lines = []
    total = len(d.paragraphs)
    for i, p in enumerate(d.paragraphs, 1):
        if progress and (i % 200 == 0 or i == total):
            progress("read Word paragraphs", i, total)
        t = (p.text or "").strip()
        if t:
            lines.append((i, t))
    return lines


def _diff_text(a_lines, b_lines, progress=None):
    a = [x[1] for x in a_lines]
    b = [x[1] for x in b_lines]
    if progress:
        progress("compare paragraphs", None, None)
    sm = difflib.SequenceMatcher(a=a, b=b)
    changes = []
    opcodes = sm.get_opcodes()
    total = len(opcodes)
    for idx, (tag, i1, i2, j1, j2) in enumerate(opcodes, 1):
        if progress and (idx % 100 == 0 or idx == total):
            progress("scan paragraph changes", idx, total)
        if tag == "equal":
            continue
        changes.append({
            "类型": tag,
            "原段落区间": f"{i1+1}-{i2}",
            "新段落区间": f"{j1+1}-{j2}",
            "原内容": "\n".join(a[i1:i2]),
            "修改后": "\n".join(b[j1:j2]),
        })
    return changes


def _write_styled_report(out: Path, df: pd.DataFrame, is_word: bool):
    if df.empty:
        if is_word:
            df = pd.DataFrame([{"类型":"无差异","原段落区间":"","新段落区间":"","原内容":"","修改后":""}])
        else:
            df = pd.DataFrame([{"Sheet":"","Cell":"","Old":"","New":"无差异"}])

    with pd.ExcelWriter(out, engine="openpyxl") as w:
        df.to_excel(w, index=False, sheet_name="差异报告")

    wb = load_workbook(out)
    ws = wb["差异报告"]
    ws.freeze_panes = "A2"

    header_fill = PatternFill("solid", fgColor="1F4E79")
    header_font = Font(color="FFFFFF", bold=True)

    for c in range(1, ws.max_column + 1):
        cell = ws.cell(1, c)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        if is_word:
            ws.column_dimensions[get_column_letter(c)].width = 24 if c < 4 else 60
        else:
            ws.column_dimensions[get_column_letter(c)].width = 18 if c <= 2 else 60
    wb.save(out)


@bridge.expose
def run_diff(a_path, b_path, strict):
    state = {"stage": "prepare", "current": None, "total": None}
    stop_event, heartbeat_thread = _start_heartbeat("diff", state)
    try:
        def progress(stage, current=None, total=None):
            state["stage"] = stage
            state["current"] = current
            state["total"] = total

        a, b = Path(a_path), Path(b_path)
        if a.suffix.lower() != b.suffix.lower():
            return {"status": "error", "msg": "两个文件类型必须一致（同为 docx 或 xlsx）"}

        ext = a.suffix.lower()
        if ext == ".docx":
            if docx is None:
                return {"status": "error", "msg": "未安装 python-docx 依赖"}
            bridge.update_terminal("[*] 正在进行 Word 段落深度比对...")
            progress(f"read {a.name}")
            a_lines = _read_docx_lines(a, progress)
            progress(f"read {b.name}")
            b_lines = _read_docx_lines(b, progress)
            changes = _diff_text(a_lines, b_lines, progress)
            out = b.parent / f"{b.stem}_Word对比报告.xlsx"
            progress("write report")
            _write_styled_report(out, pd.DataFrame(changes), is_word=True)
            return {"status": "success", "msg": f"报告已生成: {out.name}"}

        elif ext in (".xlsx", ".xls"):
            bridge.update_terminal("[*] 正在进行 Excel 逐单元格比对...")
            progress(f"open {a.name}")
            wa = load_workbook(a, data_only=True)
            progress(f"open {b.name}")
            wb = load_workbook(b, data_only=True)
            changes = []
            sheets = sorted(set(wa.sheetnames) | set(wb.sheetnames))
            for sheet_idx, s in enumerate(sheets, 1):
                sa = wa[s] if s in wa.sheetnames else None
                sb = wb[s] if s in wb.sheetnames else None
                max_r = max(sa.max_row if sa else 1, sb.max_row if sb else 1)
                max_c = max(sa.max_column if sa else 1, sb.max_column if sb else 1)
                progress(f"scan sheet {s}", sheet_idx, len(sheets))
                for r in range(1, max_r + 1):
                    if r % 500 == 0 or r == max_r:
                        bridge.update_terminal(f"  └─ Sheet {s}: row {r}/{max_r}")
                    for c in range(1, max_c + 1):
                        va = sa.cell(r, c).value if sa else None
                        vb = sb.cell(r, c).value if sb else None
                        if va != vb:
                            changes.append({"Sheet": s, "Cell": f"{get_column_letter(c)}{r}", "Old": str(va) if va is not None else "", "New": str(vb) if vb is not None else ""})

            out = b.parent / f"{b.stem}_Excel对比报告.xlsx"
            progress("write report")
            _write_styled_report(out, pd.DataFrame(changes), is_word=False)
            return {"status": "success", "msg": f"报告已生成: {out.name}"}
        else:
            return {"status": "error", "msg": "仅支持 docx 或 xlsx 比对"}

    except Exception as e:
        return {"status": "error", "msg": str(e)}
    finally:
        _stop_heartbeat(stop_event, heartbeat_thread)
