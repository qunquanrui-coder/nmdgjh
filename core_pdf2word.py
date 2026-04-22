# -*- coding: utf-8 -*-
import bridge, tempfile, fitz, os, threading, time
from pathlib import Path
from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _convert_image_mode(pdf_path, docx_path, dpi, progress_callback):
    """纯图模式固化逻辑"""
    doc = fitz.open(pdf_path)
    word_doc = Document()
    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            section = word_doc.sections[0]
            section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Mm(12.7)
            aw, ah = (section.page_width - Mm(25.4)) * 0.985, (section.page_height - Mm(25.4)) * 0.985
            
            # 删除自带的空段落
            if word_doc.paragraphs: 
                p0 = word_doc.paragraphs[0]
                p0._element.getparent().remove(p0._element)
            
            total = len(doc)
            for i in range(total):
                # 每渲染 2 页发一次心跳进度到前端
                if i % 2 == 0: 
                    progress_callback(f"  └─ 正在固化图片页 {i+1}/{total}...")
                    
                pix = doc[i].get_pixmap(matrix=fitz.Matrix(int(dpi)/72, int(dpi)/72), alpha=False)
                img_p = Path(tmp_dir) / f"{i}.png"
                pix.save(str(img_p))
                
                para = word_doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i > 0: para.paragraph_format.page_break_before = True 
                
                asp = pix.width / max(1, pix.height)
                fw, fh = aw, aw / asp
                if fh > ah: fh, fw = ah, ah * asp
                para.add_run().add_picture(str(img_p), width=fw, height=fh)
                
            word_doc.save(str(docx_path))
    finally: 
        doc.close()

def _run_editable_mode(f, out, res_container):
    """在独立子线程中运行耗时极高的解析引擎，防止堵塞 Web 界面"""
    try:
        from pdf2docx import Converter
        cv = Converter(str(f))
        cv.convert(str(out))
        cv.close()
        res_container['status'] = 'ok'
    except ImportError:
        res_container['status'] = 'error'
        res_container['error'] = "缺少 pdf2docx 库，如果已打包请检查 build_modern.py 是否漏了 --hidden-import=pdf2docx"
    except Exception as e:
        res_container['status'] = 'error'
        res_container['error'] = str(e)

def _run_image_mode(f, out, dpi, res_container):
    try:
        _convert_image_mode(f, out, dpi, lambda msg: res_container.update({'msg': msg}))
        res_container['status'] = 'ok'
    except Exception as e:
        res_container['status'] = 'error'
        res_container['error'] = str(e)

@bridge.expose
def run_pdf2word(path_str, mode, dpi):
    try:
        p = Path(path_str.strip()).resolve()
        files = []
        
        # 1. 完美支持传入文件夹进行递归扫描
        if p.is_file():
            files = [p]
        elif p.is_dir():
            bridge.update_terminal(f"[*] 检测到文件夹，正在扫描 PDF 文档: {p.name}")
            files = list(p.rglob("*.pdf"))
        else:
            return {"status": "error", "msg": "选中的路径不存在或无效。"}
            
        valid_files = [f for f in files if f.suffix.lower() == ".pdf" and not f.name.startswith("~$")]
        
        if not valid_files:
            return {"status": "error", "msg": "该路径下未找到任何有效的 PDF 文件。"}

        # 2. 开始逐个转换
        for idx, f in enumerate(valid_files, 1):
            out = f.with_suffix(".docx")
            if out.exists():
                out = f.with_name(f"{f.stem}_转换_{int(time.time())}{out.suffix}")
            
            bridge.update_terminal(f"[*] 正在转换 ({idx}/{len(valid_files)}): {f.name}")
            time.sleep(0.1) # ✨ 核心修复：原生线程必须使用 time.sleep 避开协程冲突
            
            res_container = {'status': 'pending', 'msg': ''}
            
            # 3. 将高负载任务挂入独立线程
            if mode == "editable":
                bridge.update_terminal(f"  └─ 正在深度分析排版 (若文件超大可能需要几分钟，请耐心等待)...")
                t = threading.Thread(target=_run_editable_mode, args=(f, out, res_container))
            else:
                t = threading.Thread(target=_run_image_mode, args=(f, out, dpi, res_container))
            
            t.start()
            
            # 4. 看门狗心跳循环：让前端 WebSocket 保持存活，并实时打印日志
            started_at = time.time()
            last_heartbeat = 0.0
            last_msg = ""
            while t.is_alive():
                time.sleep(0.5) # keep the native worker thread responsive
                curr_msg = res_container.get('msg', '')
                if curr_msg and curr_msg != last_msg:
                    bridge.update_terminal(curr_msg)
                    last_msg = curr_msg

                now = time.time()
                if now - last_heartbeat >= 5:
                    elapsed = int(now - started_at)
                    if mode == "editable":
                        stage = "editable conversion layout analysis"
                    else:
                        stage = curr_msg or "image conversion"
                    bridge.update_terminal(f"HB [{f.name}] {stage} still running, elapsed {elapsed}s")
                    last_heartbeat = now
            
            if res_container.get('status') == 'error':
                raise RuntimeError(res_container.get('error'))
                
            bridge.update_terminal(f"  └─ ✅ 成功保存至同目录: {out.name}")

        return {"status": "success", "msg": f"共计 {len(valid_files)} 个文件转换完毕！"}
    except Exception as e: 
        import traceback
        traceback.print_exc()
        return {"status": "error", "msg": str(e)}
